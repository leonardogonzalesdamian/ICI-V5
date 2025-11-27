from io import BytesIO
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt


def _agregar_titulo(doc: Document, texto: str, size: int = 16, negrita: bool = True):
    """
    Utilidad sencilla para añadir títulos al documento.
    """
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrita
    run.font.size = Pt(size)


def _agregar_parrafo(doc: Document, texto: str, size: int = 11, negrita: bool = False):
    """
    Añade un párrafo de texto normal.
    """
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrita
    run.font.size = Pt(size)


def _agregar_tabla_criterios(doc: Document, criterios: Dict[str, int]):
    """
    Construye una tabla con los criterios C1–C12 y sus puntajes.
    """
    if not criterios:
        _agregar_parrafo(doc, "No se encontraron criterios para mostrar.")
        return

    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"

    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = "Criterio"
    hdr_cells[1].text = "Puntaje (0–100)"

    for clave, valor in criterios.items():
        row_cells = tabla.add_row().cells
        row_cells[0].text = str(clave)
        row_cells[1].text = str(valor)


def _agregar_seccion_incongruencias(doc: Document, incong: Any):
    """
    Escribe las incongruencias encontradas, si las hay.
    `incong` puede ser lista, dict o texto.
    """
    if not incong:
        _agregar_parrafo(doc, "No se registran incongruencias específicas reportadas por el sistema.")
        return

    if isinstance(incong, str):
        _agregar_parrafo(doc, incong)
    elif isinstance(incong, list):
        for i, item in enumerate(incong, start=1):
            _agregar_parrafo(doc, f"{i}. {item}")
    elif isinstance(incong, dict):
        for clave, valor in incong.items():
            _agregar_parrafo(doc, f"- {clave}: {valor}")
    else:
        _agregar_parrafo(doc, str(incong))


def generar_informe(texto: str, resultados: Dict[str, Any], incong: Any) -> bytes:
    """
    FUNCIÓN PRINCIPAL que usa la app de Streamlit.

    Recibe:
    - texto: texto completo de la sentencia analizada.
    - resultados: diccionario devuelto por `evaluar_todo`, con:
        {
            "criterios": { "C1": ..., ..., "C12": ... },
            "ICI_sin_penalizacion": ...,
            "ICI_ajustado": ...,
            "interpretacion": ...
        }
    - incong: resu
